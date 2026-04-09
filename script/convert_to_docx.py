#!/usr/bin/env python3
"""将Markdown剧本转换为DOCX格式"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取markdown文件
with open('/Users/leifu/Movies/demo1/短剧/天工开物/短剧剧本_天工开物_36集_展开版.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建Word文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 处理每一行
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()

    # 跳过空行
    if not line:
        i += 1
        continue

    # 标题处理
    if line.startswith('# '):
        p = doc.add_heading(line[2:], level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=2)
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=3)

    # 场景标题 **【...】**
    elif line.startswith('**【') and line.endswith('】**'):
        p = doc.add_paragraph()
        run = p.add_run(line.replace('**', ''))
        run.bold = True
        run.font.size = Pt(12)

    # 场景描述 ```...```
    elif line.startswith('```') and line.endswith('```'):
        # 跳过代码块标记，添加场景描述
        desc = line[3:-3].strip()
        if desc:
            p = doc.add_paragraph(desc)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

    # 普通对话 A：（B）C
    elif re.match(r'^[\u4e00-\u9fa5a-zA-Z0-9\uff00-\uffff（）\(\)【】\[\]《》〈〉『』「」]+：', line):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)

        # 分割说话人 和 内容
        match = re.match(r'^([^：]+)：(.+)$', line)
        if match:
            speaker = match.group(1)
            speech = match.group(2)

            # 说话人加粗
            run_speaker = p.add_run(speaker + '：')
            run_speaker.bold = True

            # 内容
            p.add_run(speech)

    # 旁白/内心独白
    elif line.startswith('旁白：') or line.startswith('内心独白：') or line.startswith('画面：'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(line)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

    # 列表项
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)

    # 表格行（简化处理）
    elif line.startswith('|'):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.3)

    # 分隔线
    elif line == '---':
        p = doc.add_paragraph('─' * 40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 其他段落
    else:
        p = doc.add_paragraph(line)

    i += 1

# 保存
output_path = '/Users/leifu/Movies/demo1/短剧/天工开物/短剧剧本_天工开物_36集_展开版.docx'
doc.save(output_path)
print(f'已保存到: {output_path}')
